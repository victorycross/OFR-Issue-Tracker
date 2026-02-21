#!/usr/bin/env python3
"""
Generate the One Firm Risk Tracker User Guide as a Word (.docx) document.
Mirrors the HTML user guide content and embeds all 13 screenshots.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS = os.path.join(BASE, "screenshots")
OUTPUT = os.path.join(BASE, "User_Guide_One_Firm_Risk_Tracker.docx")

# ── Colors ──
NAVY_950  = RGBColor(0x0A, 0x16, 0x28)
NAVY_900  = RGBColor(0x0D, 0x1F, 0x3C)
NAVY_800  = RGBColor(0x16, 0x2D, 0x50)
NAVY_700  = RGBColor(0x1E, 0x3A, 0x5F)
NAVY_300  = RGBColor(0x7D, 0xA3, 0xC5)
PWC_ORANGE = RGBColor(0xD0, 0x4A, 0x02)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
GRAY_600  = RGBColor(0x64, 0x74, 0x8B)
GRAY_400  = RGBColor(0x94, 0xA3, 0xB8)
GREEN     = RGBColor(0x1E, 0x8E, 0x3E)
AMBER     = RGBColor(0xE8, 0x71, 0x0A)
RED       = RGBColor(0xD9, 0x30, 0x25)
BLUE      = RGBColor(0x25, 0x63, 0xEB)


def set_cell_shading(cell, color_hex):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_border_bottom(paragraph, color="D04A02", size="12"):
    """Add a colored bottom border to a paragraph."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def set_paragraph_spacing(paragraph, before=0, after=0, line=None):
    """Set paragraph spacing in points."""
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)


def add_run(paragraph, text, bold=False, italic=False, size=11, color=None, font_name="Calibri"):
    """Add a styled run to a paragraph."""
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font_name
    if color:
        run.font.color.rgb = color
    return run


def add_screenshot(doc, filename, caption, width=6.2):
    """Insert a screenshot image with a caption below it."""
    filepath = os.path.join(SCREENSHOTS, filename)
    if not os.path.exists(filepath):
        p = doc.add_paragraph()
        add_run(p, f"[Screenshot not found: {filename}]", italic=True, color=RED, size=10)
        return

    # Add the image
    doc.add_picture(filepath, width=Inches(width))
    # Center the last paragraph (the image)
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Caption
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(cap, before=4, after=16)
    add_run(cap, f"Figure: {caption}", italic=True, size=9, color=GRAY_600)


def add_callout(doc, prefix, text, prefix_color=GREEN):
    """Add a styled callout/tip paragraph."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=8, after=12)
    # Add left indent to simulate callout box
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8)
    add_run(p, f"{prefix}: ", bold=True, size=10, color=prefix_color)
    add_run(p, text, size=10, color=GRAY_600)


def add_step(doc, number, text_parts):
    """Add a numbered step. text_parts is list of (text, bold) tuples."""
    p = doc.add_paragraph(style="List Number")
    set_paragraph_spacing(p, before=4, after=4)
    for text, bold in text_parts:
        add_run(p, text, bold=bold, size=11)


def add_bullet(doc, text_parts):
    """Add a bullet point. text_parts is list of (text, bold) tuples."""
    p = doc.add_paragraph(style="List Bullet")
    set_paragraph_spacing(p, before=2, after=2)
    for text, bold in text_parts:
        add_run(p, text, bold=bold, size=11)


def add_section_header(doc, number, title):
    """Add a numbered section heading with orange accent."""
    # Add some space
    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=0)

    h = doc.add_heading(f"{number}. {title}" if number else title, level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY_900
        run.font.size = Pt(22)
        run.font.name = "Calibri"
    add_border_bottom(h)
    set_paragraph_spacing(h, before=24, after=12)


def add_subsection_header(doc, title):
    """Add a subsection heading."""
    h = doc.add_heading(title, level=2)
    for run in h.runs:
        run.font.color.rgb = NAVY_800
        run.font.size = Pt(15)
        run.font.name = "Calibri"
    set_paragraph_spacing(h, before=16, after=8)


def add_body(doc, text, color_override=None):
    """Add a body text paragraph."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=8)
    add_run(p, text, size=11, color=color_override or RGBColor(0x1E, 0x29, 0x3B))
    return p


def add_lead(doc, text):
    """Add a lead/intro paragraph in muted color."""
    p = doc.add_paragraph()
    set_paragraph_spacing(p, before=4, after=14)
    add_run(p, text, size=11.5, color=GRAY_600)


def create_cover_page(doc):
    """Create the title/cover page."""
    # Blank spacing at top
    for _ in range(4):
        spacer = doc.add_paragraph()
        set_paragraph_spacing(spacer, before=0, after=0)

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(title, "One Firm", bold=True, size=36, color=NAVY_900, font_name="Calibri")
    add_run(title, " Risk Tracker", bold=True, size=36, color=PWC_ORANGE, font_name="Calibri")
    set_paragraph_spacing(title, before=0, after=4)

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(sub, "User Guide", bold=False, size=24, color=NAVY_700, font_name="Calibri")
    set_paragraph_spacing(sub, before=0, after=24)

    # Orange rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(rule, "\u2500" * 40, size=12, color=PWC_ORANGE)
    set_paragraph_spacing(rule, before=0, after=24)

    # Description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(desc, before=0, after=32)
    add_run(desc, "A modern dashboard for tracking risks and issues\nacross the One Firm Risk group.", size=13, color=GRAY_600)

    # Meta info table
    meta_table = doc.add_table(rows=1, cols=3)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cells = meta_table.rows[0].cells
    for i, (label, value) in enumerate([
        ("Platform", "Web Browser"),
        ("Authentication", "PwC Microsoft SSO"),
        ("Languages", "English / French")
    ]):
        p = cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, f"{label}\n", bold=True, size=9, color=NAVY_700)
        add_run(p, value, size=10, color=GRAY_600)

    # Spacing
    for _ in range(6):
        spacer = doc.add_paragraph()
        set_paragraph_spacing(spacer, before=0, after=0)

    # Footer info
    footer_line = doc.add_paragraph()
    footer_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(footer_line, "Version 1.0  |  February 2026", size=10, color=GRAY_400)
    set_paragraph_spacing(footer_line, before=0, after=4)

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_run(org, "PwC Canada  \u2014  OFR Risk & Quality Team", bold=True, size=10, color=NAVY_700)

    # Page break
    doc.add_page_break()


def create_toc_page(doc):
    """Create a table of contents page."""
    h = doc.add_heading("Table of Contents", level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY_900
        run.font.size = Pt(22)
        run.font.name = "Calibri"
    add_border_bottom(h)
    set_paragraph_spacing(h, before=12, after=20)

    toc_items = [
        ("1", "Getting Started"),
        ("2", "Weekly Review Dashboard"),
        ("3", "Submitting New Issues"),
        ("4", "Risk & Issue Register"),
        ("5", "Adding Updates"),
        ("6", "Filtering & Searching"),
        ("7", "Exporting to CSV"),
        ("8", "Bilingual Support (EN / FR)"),
        ("9", "Putting It All Together"),
        ("", "Frequently Asked Questions"),
    ]
    for num, title in toc_items:
        p = doc.add_paragraph()
        set_paragraph_spacing(p, before=6, after=6)
        if num:
            add_run(p, f"  {num}.  ", bold=True, size=11, color=PWC_ORANGE)
        else:
            add_run(p, "       ", size=11)
        add_run(p, title, size=11, color=NAVY_800)

    doc.add_page_break()


def create_feature_table(doc, features):
    """Create a 2-column feature description table."""
    cols = 2
    rows_needed = (len(features) + 1) // 2
    table = doc.add_table(rows=rows_needed, cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for idx, (title, desc) in enumerate(features):
        row = idx // cols
        col = idx % cols
        cell = table.rows[row].cells[col]
        set_cell_shading(cell, "F8FAFC")

        p = cell.paragraphs[0]
        add_run(p, title, bold=True, size=10, color=NAVY_800)
        p2 = cell.add_paragraph()
        set_paragraph_spacing(p2, before=2, after=4)
        add_run(p2, desc, size=9.5, color=GRAY_600)

    # Style borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=8)


# ══════════════════════════════════════════════════
#  MAIN
# ══════════════════════════════════════════════════
def main():
    doc = Document()

    # ── Page setup ──
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # ── Default font ──
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # ═══ COVER PAGE ═══
    create_cover_page(doc)

    # ═══ TABLE OF CONTENTS ═══
    create_toc_page(doc)

    # ═══ WHAT IS IT ═══
    h = doc.add_heading("What is the One Firm Risk Tracker?", level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY_900
        run.font.size = Pt(22)
        run.font.name = "Calibri"
    add_border_bottom(h)
    set_paragraph_spacing(h, before=12, after=12)

    add_lead(doc, "The One Firm Risk Tracker is an interactive web dashboard that gives the OFR team a single, shared view of all open issues with built-in staleness tracking, update history, and weekly review metrics. It drives the full issue lifecycle \u2014 from identification through active management to resolution.")

    create_feature_table(doc, [
        ("Weekly Review Dashboard", "At-a-glance KPI cards showing open items, stale items, and priority breakdown."),
        ("Issue Intake", "Submit new risks directly into a triage queue for committee review."),
        ("Update Tracking", "Add timestamped updates with full audit history on every item."),
        ("CSV Export", "Export all data to CSV for offline use or SharePoint List import."),
        ("Staleness Alerts", "Color-coded indicators flag items that haven't been updated recently."),
        ("Bilingual (EN/FR)", "Full English and French language support with one-click toggle."),
    ])

    doc.add_page_break()

    # ═══ 1. GETTING STARTED ═══
    add_section_header(doc, "1", "Getting Started")
    add_lead(doc, "Open the Risk Tracker in your browser. You will see the sign-in screen with two options: sign in with your PwC account, or try the demo.")

    add_screenshot(doc, "01_login_screen.png", "The login screen with PwC Microsoft sign-in and Demo Mode access")

    add_subsection_header(doc, "How to sign in")

    add_step(doc, 1, [
        ('Click ', False), ('"Sign in with Microsoft"', True),
        (' to authenticate with your PwC account. You may be prompted for multi-factor authentication (MFA).', False)
    ])
    add_step(doc, 2, [
        ('Once signed in, click ', False), ('"Connect to SharePoint"', True),
        (' to load your team\'s live risk data. A green ', False), ('Connected', True),
        (' badge will appear in the header.', False)
    ])
    add_step(doc, 3, [
        ('Alternatively, click ', False), ('"Try Demo"', True),
        (' to explore the tool with sample data \u2014 no sign-in required. An amber ', False),
        ('Demo Mode', True), (' badge indicates you\'re using sample data.', False)
    ])

    add_screenshot(doc, "13_header_demo_mode.png", "The header bar displays your connection status and language toggle", width=5.0)

    add_callout(doc, "Tip", "You can switch between Demo Mode and Live Mode at any time by signing in or out. Demo Mode is useful for training new team members without affecting real data.", GREEN)

    doc.add_page_break()

    # ═══ 2. DASHBOARD ═══
    add_section_header(doc, "2", "Weekly Review Dashboard")
    add_lead(doc, "The top of the screen shows a real-time summary at a glance \u2014 designed for the weekly committee review.")

    add_screenshot(doc, "03_kpi_cards.png", "The dashboard with KPI summary cards, intake panel, and issue tracker")

    add_subsection_header(doc, "What the numbers mean")

    create_feature_table(doc, [
        ("Open Items", "Total number of risks/issues currently being tracked (excludes Closed items)."),
        ("Stale", "Items that haven't been updated in 14+ days. Shown in red to flag attention."),
        ("High Priority", "Number of items marked as High priority that are still open."),
        ("Medium / Low", "Breakdown of remaining items by priority level."),
    ])

    add_body(doc, "Below the KPI cards, you will see a status breakdown showing how many items are Active, Monitoring, Escalated, or New.")

    doc.add_page_break()

    # ═══ 3. INTAKE ═══
    add_section_header(doc, "3", "Submitting New Issues")
    add_lead(doc, "The Issue Intake panel is where new risks and issues enter the system. Items submitted here go into a triage queue for committee review before being added to the formal register.")

    add_screenshot(doc, "04_intake_panel.png", "The Issue Intake section with pending items shown as cards")

    add_subsection_header(doc, "Submitting a new issue")

    add_step(doc, 1, [
        ('Click the ', False), ('"New Issue"', True),
        (' button in the top-right of the Intake section.', False)
    ])
    add_step(doc, 2, [
        ('Fill in the ', False), ('title', True), (', ', False), ('owner', True),
        (' (who is responsible), ', False), ('priority', True),
        (' (High, Medium, or Low), and an optional ', False), ('description', True), ('.', False)
    ])
    add_step(doc, 3, [
        ('Click ', False), ('Submit', True),
        ('. The item appears as a card in the intake queue.', False)
    ])

    add_subsection_header(doc, "What happens next")
    add_body(doc, "Each intake card has two actions:")

    add_bullet(doc, [
        ("Move to Tracker", True),
        (" \u2014 promotes the item into the formal issue tracker. It gets a tracking ID (e.g., OFR-009) and enters the active lifecycle.", False)
    ])
    add_bullet(doc, [
        ("Dismiss (X)", True),
        (" \u2014 removes the item from the queue if it's not relevant or a duplicate.", False)
    ])

    add_callout(doc, "Note", "Only committee members with the appropriate permissions can promote or dismiss intake items. Regular users can submit issues for review.", BLUE)

    doc.add_page_break()

    # ═══ 4. RISK REGISTER ═══
    add_section_header(doc, "4", "Issue Tracker")
    add_lead(doc, "The main table is the heart of the tool \u2014 a sortable, filterable view of all tracked issues with real-time staleness indicators and full lifecycle history.")

    add_screenshot(doc, "05_risk_register_table.png", "The Issue Tracker showing all actively managed items")

    add_subsection_header(doc, "Understanding the columns")

    create_feature_table(doc, [
        ("ID", "Unique tracking number (OFR-001, OFR-002, etc.) assigned when an item enters the register."),
        ("Topic", "Brief description of the risk or issue being tracked."),
        ("Owner", "The person responsible for managing and updating this item."),
        ("Priority", "Severity level: High, Medium, or Low."),
        ("Status", "Current state: Active, New, Escalated, or Monitoring."),
        ("Days Since Update", "Color-coded staleness indicator. Green = recent, amber = aging, red = overdue."),
    ])

    add_subsection_header(doc, "Staleness indicators")
    add_body(doc, "The Days Since Update column uses color coding to flag items that need attention:")

    add_screenshot(doc, "12_staleness_indicators.png", "Color-coded staleness: green (recent), amber (aging), red (overdue)")

    # Staleness legend table
    legend_table = doc.add_table(rows=3, cols=2)
    legend_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row_idx, (color_hex, label, desc) in enumerate([
        ("DCFCE7", "Green (0\u20136 days)", "Recently updated"),
        ("FEF3C7", "Amber (7\u201313 days)", "Aging, consider updating"),
        ("FEE2E2", "Red (14+ days)", "Stale, update required"),
    ]):
        cell_color = legend_table.rows[row_idx].cells[0]
        cell_desc = legend_table.rows[row_idx].cells[1]
        set_cell_shading(cell_color, color_hex)
        p1 = cell_color.paragraphs[0]
        add_run(p1, label, bold=True, size=10, color=NAVY_800)
        p2 = cell_desc.paragraphs[0]
        add_run(p2, desc, size=10, color=GRAY_600)

    spacer = doc.add_paragraph()
    set_paragraph_spacing(spacer, before=0, after=12)

    add_subsection_header(doc, "Expanding a row")
    add_body(doc, "Click the arrow icon on the left side of any row to expand it. The expanded view shows the complete Update History \u2014 every status change and note that has been recorded for that item, displayed as a timeline.")

    add_screenshot(doc, "06_expanded_row.png", "An expanded row revealing the full update history with timestamps and status changes")

    doc.add_page_break()

    # ═══ 5. UPDATES ═══
    add_section_header(doc, "5", "Adding Updates")
    add_lead(doc, "Keeping items up to date is essential. The update workflow lets you add notes and optionally change an item's status in one step.")

    add_screenshot(doc, "07_update_modal.png", "The Add Update dialog with notes field and optional status change")

    add_subsection_header(doc, "How to add an update")

    add_step(doc, 1, [
        ("Expand the row", True),
        (" by clicking the arrow icon on the left side of the item you want to update.", False)
    ])
    add_step(doc, 2, [
        ("Click the ", False), ("pencil icon", True),
        (" (Add Update) that appears in the expanded section. This opens the update dialog.", False)
    ])
    add_step(doc, 3, [
        ("Type your ", False), ("update notes", True),
        (" in the text area. Be specific about what has changed, what actions were taken, or what is planned next.", False)
    ])
    add_step(doc, 4, [
        ("(Optional) Use the ", False), ("Change Status", True),
        (' dropdown to move the item to a different status \u2014 for example, from "Active" to "Monitoring" or "Escalated".', False)
    ])
    add_step(doc, 5, [
        ("Click ", False), ('"Save Update"', True),
        (". The note is timestamped and added to the item's history. The \"Days Since Update\" counter resets to ", False),
        ("Today", True), (".", False)
    ])

    add_callout(doc, "Tip", "During the weekly review meeting, go through each stale item (red indicators) and add updates to keep the register current. This ensures the dashboard KPIs accurately reflect the team's risk posture.", GREEN)

    doc.add_page_break()

    # ═══ 6. FILTERS ═══
    add_section_header(doc, "6", "Filtering & Searching")
    add_lead(doc, "Use the filter bar above the register table to quickly find the items you need.")

    add_screenshot(doc, "08_filters_and_controls.png", "Filter buttons for quick views, plus the search bar for text-based filtering")

    add_subsection_header(doc, "Available filters")

    add_bullet(doc, [("All Open", True), (" \u2014 Shows all non-closed items (default view)", False)])
    add_bullet(doc, [("Stale Items", True), (" \u2014 Shows only items not updated in 14+ days (the count is shown on the button)", False)])
    add_bullet(doc, [("My Items", True), (" \u2014 Shows only items assigned to you (uses your signed-in name in Live Mode)", False)])
    add_bullet(doc, [("High / Medium / Low Priority", True), (" \u2014 Filter by priority level", False)])

    add_subsection_header(doc, "Searching")
    add_body(doc, "The search box next to the filter buttons lets you type any keyword to search across all item topics. Results update instantly as you type.")

    doc.add_page_break()

    # ═══ 7. EXPORT ═══
    add_section_header(doc, "7", "Exporting to CSV")
    add_lead(doc, "Export the full issue tracker to a CSV file for reporting, offline use, or importing into a SharePoint List.")

    add_screenshot(doc, "09_csv_export.png", "The Export CSV button in the filter bar area")

    add_subsection_header(doc, "How to export")

    add_step(doc, 1, [
        ("Click the ", False), ('"Export CSV"', True),
        (" button in the top-right area of the Issue Tracker section.", False)
    ])
    add_step(doc, 2, [
        ("A CSV file named ", False), ("OFR_Risk_Register_YYYY-MM-DD.csv", True),
        (" downloads automatically.", False)
    ])
    add_step(doc, 3, [
        ("Open in Excel, Google Sheets, or import directly into a SharePoint List.", False)
    ])

    add_callout(doc, "Note", "The CSV export includes all fields and all update history in a denormalized format (one row per update). It uses UTF-8 encoding with BOM for proper character display in Excel, including French accented characters.", BLUE)

    doc.add_page_break()

    # ═══ 8. BILINGUAL ═══
    add_section_header(doc, "8", "Bilingual Support (EN / FR)")
    add_lead(doc, "The entire application is available in both English and French, reflecting PwC Canada's bilingual requirements.")

    add_screenshot(doc, "10_french_language.png", "The complete interface displayed in French after toggling the language")

    add_subsection_header(doc, "Switching languages")
    add_body(doc, "Click the FR button in the top-right corner of the header to switch to French. Click EN to switch back to English. The toggle is available on every screen, including the login page.")

    add_body(doc, "When in French mode:")
    add_bullet(doc, [("All labels, buttons, and headings are translated", False)])
    add_bullet(doc, [('KPI card titles update (e.g., "Open Items" becomes "Elements ouverts")', False)])
    add_bullet(doc, [('Status labels translate (e.g., "Active" becomes "Actif")', False)])
    add_bullet(doc, [("The footer and system messages also switch language", False)])

    add_callout(doc, "Note", "Your data (risk topics, owner names, update notes) is not translated \u2014 it stays in whatever language it was entered in. Only the application interface switches languages.", BLUE)

    doc.add_page_break()

    # ═══ 9. FULL VIEW ═══
    add_section_header(doc, "9", "Putting It All Together")
    add_lead(doc, "Here is the full application as it appears during a typical weekly review session.")

    add_screenshot(doc, "02_dashboard_overview.png", "Complete application view: dashboard KPIs, intake panel, and issue tracker in one scrollable page")

    doc.add_page_break()

    # ═══ FAQ ═══
    h = doc.add_heading("Frequently Asked Questions", level=1)
    for run in h.runs:
        run.font.color.rgb = NAVY_900
        run.font.size = Pt(22)
        run.font.name = "Calibri"
    add_border_bottom(h)
    set_paragraph_spacing(h, before=12, after=16)

    faq = [
        ("Where is my data stored?",
         "In Live Mode, all data is stored in SharePoint Online Lists within the PwC M365 tenant. In Demo Mode, data is temporary and stored only in your browser session \u2014 it resets when you close the tab."),
        ("Who can see the data?",
         "Only users with access to the SharePoint site can view or edit risk data. Access is managed by the OFR team through standard SharePoint permissions."),
        ("Can I use this on my phone?",
         "The application is responsive and works on tablets and phones, but it is designed primarily for desktop use during weekly review meetings."),
        ("What if SharePoint is down?",
         "The application automatically falls back to Demo Mode if it cannot connect to SharePoint. You will see an error banner with options to retry or switch to demo data. No data is lost \u2014 SharePoint has its own redundancy."),
        ("How do I get access?",
         "Submit a request through ServiceNow to be added to the OFR Risk Tracker SharePoint site. Your request will be approved by the OFR Risk Lead."),
        ("Who do I contact for help?",
         "Reach out to the OFR Risk & Quality team via the #ofr-risk-tracker Teams channel, or email the OFR support mailbox."),
    ]

    for q, a in faq:
        qp = doc.add_paragraph()
        set_paragraph_spacing(qp, before=12, after=4)
        add_run(qp, q, bold=True, size=12, color=NAVY_800)

        ap = doc.add_paragraph()
        set_paragraph_spacing(ap, before=2, after=8)
        add_run(ap, a, size=11, color=GRAY_600)

    # ═══ FOOTER ═══
    doc.add_page_break()

    footer_p1 = doc.add_paragraph()
    footer_p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer_p1, before=48, after=8)
    add_run(footer_p1, "One Firm Risk Tracker \u2014 User Guide v1.0 \u2014 February 2026", size=11, color=GRAY_400)

    footer_p2 = doc.add_paragraph()
    footer_p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer_p2, before=4, after=8)
    add_run(footer_p2, "PwC Canada \u2014 OFR Risk & Quality Team", bold=True, size=11, color=NAVY_700)

    footer_p3 = doc.add_paragraph()
    footer_p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(footer_p3, before=4, after=0)
    add_run(footer_p3, "For related documents, see the Business Requirements Document, Executive Outline, and Security Design Document.", size=10, color=GRAY_400)

    # ── Save ──
    doc.save(OUTPUT)
    print(f"\nUser Guide DOCX created: {OUTPUT}")
    print(f"File size: {os.path.getsize(OUTPUT):,} bytes")


if __name__ == "__main__":
    main()
