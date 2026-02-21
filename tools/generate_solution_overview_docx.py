#!/usr/bin/env python3
"""
Generate the One Firm Risk Tracker — Solution Overview DOCX.
Positions the tool as an issue lifecycle management platform,
not another risk register. Focuses on value, workflow, and outcomes.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

BASE = os.path.dirname(os.path.abspath(__file__))
SCREENSHOTS = os.path.join(BASE, "screenshots")
OUTPUT = os.path.join(BASE, "Solution_Overview_One_Firm_Risk_Tracker.docx")

# ── Colors ──
NAVY_950  = RGBColor(0x0A, 0x16, 0x28)
NAVY_900  = RGBColor(0x0D, 0x1F, 0x3C)
NAVY_800  = RGBColor(0x16, 0x2D, 0x50)
NAVY_700  = RGBColor(0x1E, 0x3A, 0x5F)
NAVY_300  = RGBColor(0x7D, 0xA3, 0xC5)
PWC_ORANGE = RGBColor(0xD0, 0x4A, 0x02)
WHITE     = RGBColor(0xFF, 0xFF, 0xFF)
TEXT_MAIN = RGBColor(0x1E, 0x29, 0x3B)
GRAY_600  = RGBColor(0x64, 0x74, 0x8B)
GRAY_500  = RGBColor(0x78, 0x84, 0x96)
GRAY_400  = RGBColor(0x94, 0xA3, 0xB8)
GREEN     = RGBColor(0x1E, 0x8E, 0x3E)
AMBER     = RGBColor(0xE8, 0x71, 0x0A)
RED       = RGBColor(0xD9, 0x30, 0x25)
BLUE      = RGBColor(0x25, 0x63, 0xEB)
PURPLE    = RGBColor(0x6B, 0x21, 0xA8)


# ══════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

def style_table_borders(table, color="E2E8F0"):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="{color}"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def add_border_bottom(paragraph, color="D04A02", size="12"):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="{size}" w:space="4" w:color="{color}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)

def sp(paragraph, before=0, after=0, line=None):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def run(paragraph, text, bold=False, italic=False, size=11, color=None, font="Calibri"):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.name = font
    if color:
        r.font.color.rgb = color
    return r

def img(doc, filename, caption, width=6.2):
    filepath = os.path.join(SCREENSHOTS, filename)
    if not os.path.exists(filepath):
        p = doc.add_paragraph()
        run(p, f"[Image not found: {filename}]", italic=True, color=RED, size=10)
        return
    doc.add_picture(filepath, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(cap, before=4, after=16)
    run(cap, caption, italic=True, size=9, color=GRAY_600)

def body(doc, text, color_override=None):
    p = doc.add_paragraph()
    sp(p, before=4, after=10)
    run(p, text, size=11, color=color_override or TEXT_MAIN)
    return p

def lead(doc, text):
    p = doc.add_paragraph()
    sp(p, before=4, after=14)
    run(p, text, size=12, color=GRAY_500)

def h1(doc, text, numbered=True):
    spacer = doc.add_paragraph()
    sp(spacer, before=0, after=0)
    h = doc.add_heading(text, level=1)
    for r in h.runs:
        r.font.color.rgb = NAVY_900
        r.font.size = Pt(22)
        r.font.name = "Calibri"
    add_border_bottom(h)
    sp(h, before=20, after=12)

def h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.color.rgb = NAVY_800
        r.font.size = Pt(15)
        r.font.name = "Calibri"
    sp(h, before=16, after=8)

def bullet(doc, text_parts):
    p = doc.add_paragraph(style="List Bullet")
    sp(p, before=2, after=2)
    for text, bold in text_parts:
        run(p, text, bold=bold, size=11)

def callout(doc, prefix, text, prefix_color=GREEN):
    p = doc.add_paragraph()
    sp(p, before=8, after=12)
    pf = p.paragraph_format
    pf.left_indent = Cm(0.8)
    run(p, f"{prefix}: ", bold=True, size=10, color=prefix_color)
    run(p, text, size=10, color=GRAY_600)


# ══════════════════════════════════════════════════
#  DOCUMENT
# ══════════════════════════════════════════════════

def main():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = TEXT_MAIN

    # ═══════════════════════════════════════════════
    #  COVER PAGE
    # ═══════════════════════════════════════════════
    for _ in range(5):
        s = doc.add_paragraph(); sp(s, 0, 0)

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(t, "One Firm", bold=True, size=38, color=NAVY_900)
    run(t, " Risk Tracker", bold=True, size=38, color=PWC_ORANGE)
    sp(t, 0, 4)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(sub, "Solution Overview", size=24, color=NAVY_700)
    sp(sub, 0, 6)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(tagline, "From identification to resolution \u2014 a single view\nthat keeps issues moving forward.", size=13, color=GRAY_500, font="Calibri")
    sp(tagline, 8, 32)

    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(rule, "\u2500" * 40, size=11, color=PWC_ORANGE)
    sp(rule, 0, 32)

    # Cover meta
    meta = doc.add_table(rows=1, cols=3)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    remove_table_borders(meta)
    for i, (lbl, val) in enumerate([
        ("Audience", "One Firm Risk Leadership"),
        ("Classification", "Internal Use"),
        ("Version", "1.0 \u2014 February 2026"),
    ]):
        p = meta.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run(p, f"{lbl}\n", bold=True, size=9, color=NAVY_700)
        run(p, val, size=10, color=GRAY_600)

    for _ in range(6):
        s = doc.add_paragraph(); sp(s, 0, 0)

    org = doc.add_paragraph()
    org.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(org, "PwC Canada  \u2014  OFR Risk & Quality", bold=True, size=10, color=NAVY_700)
    sp(org, 0, 0)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    #  THE PROBLEM
    # ═══════════════════════════════════════════════
    h1(doc, "The Challenge")
    lead(doc, "Every risk function faces the same friction: issues are identified, documented, and then slowly lose momentum. The problem has never been awareness \u2014 it has been follow-through.")

    body(doc, "Across advisory and assurance teams, the pattern is familiar. A risk or issue surfaces, it gets noted in a shared document or spreadsheet, ownership is assigned, and then the trail goes cold. Weeks later, a committee review reveals that half the register has not been touched. Not because people are not working on the issues \u2014 but because nothing is prompting them to close the loop.")

    body(doc, "The consequences compound quietly:")
    bullet(doc, [("Items go stale.", True), (" Without a visible clock on each issue, aging items blend into the background. What was urgent last month becomes furniture.", False)])
    bullet(doc, [("Accountability diffuses.", True), (" When there is no shared, real-time view of who owns what and when it was last touched, it is easy for items to sit in limbo between teams.", False)])
    bullet(doc, [("Committee time is wasted.", True), (" Weekly reviews spend more time asking \"what is the status of...\" than actually making decisions. The meeting becomes a status-gathering exercise instead of a decision-making forum.", False)])
    bullet(doc, [("Reporting is manual and retrospective.", True), (" Pulling together a summary means stitching data from multiple documents. By the time the report is assembled, it is already out of date.", False)])

    body(doc, "The One Firm Risk Tracker was built to solve these specific problems \u2014 not by adding another layer of documentation, but by making issue lifecycle management a continuous, visible, and accountable process.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    #  THE APPROACH
    # ═══════════════════════════════════════════════
    h1(doc, "The Approach")
    lead(doc, "Rather than building another static register that people update before a meeting, the solution is designed around a simple principle: issues should always be moving forward, and everyone should be able to see that motion.")

    h2(doc, "A lifecycle, not a list")
    body(doc, "The tool structures every issue around a clear lifecycle with four natural stages:")

    # Lifecycle stages table
    stages = doc.add_table(rows=4, cols=2)
    stages.alignment = WD_TABLE_ALIGNMENT.LEFT
    style_table_borders(stages, "D4E2EE")
    for i, (stage, desc, color_hex) in enumerate([
        ("Identification", "An issue is raised through the intake queue. It enters a triage holding area where the committee can review, validate, and decide whether to formally track it \u2014 or dismiss it early.", "EFF6FF"),
        ("Active Management", "Once promoted, the item is assigned an owner and enters the formal register. The staleness clock starts ticking. The owner is accountable for providing regular updates.", "F0FDF4"),
        ("Monitoring & Escalation", "As issues evolve, they can be moved to monitoring (stable, watching closely) or escalated (requires senior attention). Each transition is recorded with a timestamp and note.", "FFFBEB"),
        ("Resolution", "When an issue is resolved, it is closed with a final note. The full audit trail remains accessible for future reference, lessons learned, and regulatory evidence.", "F8FAFC"),
    ]):
        cell_stage = stages.rows[i].cells[0]
        cell_desc = stages.rows[i].cells[1]
        set_cell_shading(cell_stage, color_hex)
        set_cell_shading(cell_desc, color_hex)
        set_cell_margins(cell_stage, 80, 80, 120, 120)
        set_cell_margins(cell_desc, 80, 80, 120, 120)
        p1 = cell_stage.paragraphs[0]
        run(p1, stage, bold=True, size=11, color=NAVY_800)
        p2 = cell_desc.paragraphs[0]
        run(p2, desc, size=10, color=GRAY_600)

    spacer = doc.add_paragraph(); sp(spacer, 0, 12)

    body(doc, "This is not a passive document. Every item has a visible clock. Every status change is recorded. Every owner knows their items are being watched \u2014 not by a manager, but by the system itself.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    #  HOW IT WORKS
    # ═══════════════════════════════════════════════
    h1(doc, "How It Works")
    lead(doc, "The solution is a modern web application accessible through any browser. It connects to the firm's Microsoft 365 environment, uses corporate single sign-on, and stores data in SharePoint \u2014 no new infrastructure required.")

    # ─── Weekly Review Dashboard ───
    h2(doc, "At a Glance: The Weekly Review Dashboard")
    body(doc, "When you open the tool, the first thing you see is a real-time summary designed for the weekly committee meeting. At a glance, the leadership team can see how many issues are open, how many have gone stale, and what the priority distribution looks like. No one needs to ask \"what is the current status?\" \u2014 it is right there.")

    img(doc, "03_kpi_cards.png", "The weekly review dashboard surfaces the numbers that matter: open items, staleness, and priority distribution.")

    body(doc, "The value here is speed. Instead of spending the first 15 minutes of a review gathering status, the team walks in with a shared picture and spends their time making decisions.")

    # ─── Issue Intake ───
    h2(doc, "Capture: Issue Intake")
    body(doc, "New issues enter through a dedicated intake queue \u2014 a lightweight triage area that separates \"someone raised a concern\" from \"this is a formally tracked item.\" This distinction matters. Not every concern warrants full lifecycle tracking, and the intake step gives the committee the opportunity to validate, prioritize, and assign before an item enters the formal register.")

    img(doc, "04_intake_panel.png", "The intake queue gives the committee a clear view of what is waiting to be triaged.")

    body(doc, "Items in the intake queue can be promoted into the formal register with a single action, or dismissed if they are duplicates, out of scope, or resolved at the point of entry. This keeps the register clean and intentional.")

    # ─── The Register ───
    h2(doc, "Track: The Central Register")
    body(doc, "Once an item is promoted, it enters the central register \u2014 a live, sortable table that is the single source of truth for all tracked issues. Each item carries its ID, owner, priority, status, and a staleness indicator that shows exactly how many days have passed since the last update.")

    img(doc, "05_risk_register_table.png", "Every tracked issue in one place, with ownership and staleness visible at a glance.")

    body(doc, "The staleness indicator is the most important column in the table. It applies gentle, persistent pressure: green means recently updated, amber means aging, and red means overdue. There is no hiding. Items that are not being actively managed are immediately visible to everyone in the meeting.")

    img(doc, "12_staleness_indicators.png", "Color-coded staleness removes ambiguity. Red items demand attention; green items are healthy.")

    doc.add_page_break()

    # ─── Updates ───
    h2(doc, "Progress: The Update Workflow")
    body(doc, "The update workflow is intentionally simple. Expand an item, click the update button, type a note, and optionally change the status. One action. Ten seconds. The staleness clock resets, the update is timestamped, and the full history is preserved.")

    img(doc, "07_update_modal.png", "Adding an update is a single, focused action \u2014 designed to be done in the flow of a meeting.")

    body(doc, "Every update creates a permanent audit trail. When a committee member asks \"what happened with this issue in January?\", the answer is already there \u2014 a chronological timeline of every note, every status change, and every owner action.")

    img(doc, "06_expanded_row.png", "The full update history for any item is one click away \u2014 a chronological record of every action taken.")

    body(doc, "This is where the lifecycle model pays off. The tool does not just store issues \u2014 it tracks the velocity of resolution. Items that are being actively worked show a healthy cadence of updates. Items that have stalled show silence. The difference is immediately obvious.")

    # ─── Filtering ───
    h2(doc, "Focus: Filtering and Search")
    body(doc, "During a review meeting, the committee does not need to see everything. Quick filters let the facilitator instantly narrow the view to stale items, high-priority items, or a specific person's portfolio. A keyword search covers everything else.")

    img(doc, "08_filters_and_controls.png", "One-click filters for the views that matter most during committee reviews.")

    body(doc, "The \"Stale Items\" filter is particularly useful. Click it, and the table shows only the items that have not been updated in 14 or more days \u2014 complete with a count badge. This is the fastest path to a productive meeting: start with what needs attention, work through it, move on.")

    doc.add_page_break()

    # ─── Export ───
    h2(doc, "Report: CSV Export")
    body(doc, "When data needs to leave the tool \u2014 for a board report, an audit trail, or a SharePoint List import \u2014 the full register can be exported as a CSV with a single click. The export includes every field and the complete update history in a denormalized format ready for Excel, Sheets, or Power BI.")

    img(doc, "09_csv_export.png", "One-click export for offline analysis, reporting, or SharePoint integration.")

    # ─── Bilingual ───
    h2(doc, "Inclusive: Bilingual Support")
    body(doc, "The entire interface is available in both English and French with a one-click toggle. All labels, headings, filters, status badges, and system messages switch language instantly. Data entered by users \u2014 issue descriptions, update notes \u2014 remains in its original language, which is the correct behaviour for a bilingual team where members work in both languages.")

    img(doc, "10_french_language.png", "Full French language support reflects PwC Canada's bilingual operating environment.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    #  WHAT CHANGES
    # ═══════════════════════════════════════════════
    h1(doc, "What Changes")
    lead(doc, "The value of the tool is not in the technology. It is in the behavioural shift it creates.")

    # Before / After table
    ba = doc.add_table(rows=7, cols=2)
    ba.alignment = WD_TABLE_ALIGNMENT.LEFT
    style_table_borders(ba, "D4E2EE")

    # Header row
    for ci, (label, shade) in enumerate([("Before", "FEE2E2"), ("After", "DCFCE7")]):
        cell = ba.rows[0].cells[ci]
        set_cell_shading(cell, shade)
        set_cell_margins(cell, 80, 80, 120, 120)
        p = cell.paragraphs[0]
        run(p, label, bold=True, size=11, color=NAVY_900)

    before_after = [
        ("Issues captured in Word docs and emails", "Issues enter a structured intake queue with triage workflow"),
        ("Status discovered by asking in meetings", "Status visible in real time to everyone"),
        ("No visibility into aging or staleness", "Staleness indicators provide automatic accountability"),
        ("Update history scattered across threads", "Complete audit trail on every item"),
        ("Reporting assembled manually each week", "Dashboard and CSV export always current"),
        ("Committee reviews spend time gathering info", "Reviews start with a shared picture and focus on decisions"),
    ]
    for ri, (before, after) in enumerate(before_after, start=1):
        cell_b = ba.rows[ri].cells[0]
        cell_a = ba.rows[ri].cells[1]
        set_cell_shading(cell_b, "FEFEFE")
        set_cell_shading(cell_a, "FAFFFE")
        set_cell_margins(cell_b, 70, 70, 120, 120)
        set_cell_margins(cell_a, 70, 70, 120, 120)
        pb = cell_b.paragraphs[0]
        run(pb, before, size=10, color=GRAY_600)
        pa = cell_a.paragraphs[0]
        run(pa, after, size=10, color=NAVY_800)

    spacer = doc.add_paragraph(); sp(spacer, 0, 16)

    body(doc, "The most significant change is cultural. When every issue has a visible clock and a recorded owner, the default behaviour shifts from \"I will update this when someone asks\" to \"I should update this before the meeting.\" The tool does not enforce compliance \u2014 it creates transparency that makes compliance the natural choice.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    #  THE FULL PICTURE
    # ═══════════════════════════════════════════════
    h1(doc, "The Full Picture")
    lead(doc, "Here is the complete solution as it appears during a typical weekly review session \u2014 dashboard, intake, and register in one continuous view.")

    img(doc, "02_dashboard_overview.png", "The complete solution: dashboard metrics, intake triage, and the central register in a single scrollable view.")

    body(doc, "Everything the committee needs is on one screen. No switching between tabs. No opening attachments. No asking \"can someone share their screen?\" The facilitator opens the tool, and the meeting begins.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    #  DESIGN PRINCIPLES
    # ═══════════════════════════════════════════════
    h1(doc, "Design Principles")
    lead(doc, "The decisions behind the solution are intentional. Every feature exists to serve a specific need observed in the current process.")

    principles = [
        ("Visibility over documentation",
         "The primary output is not a report \u2014 it is a live, shared view. When the committee opens the tool, they see the current state of every issue. Documentation happens as a byproduct of using the tool, not as a separate activity."),
        ("Gentle accountability",
         "Staleness indicators create accountability without confrontation. The system does not send angry emails or generate escalation tickets. It simply makes the passage of time visible. An item that was last updated 23 days ago speaks for itself."),
        ("Low friction updates",
         "Adding an update takes less than 15 seconds. This is by design. If the update process is burdensome, people will not do it. If it is effortless, they will do it in the natural flow of their work."),
        ("Triage before tracking",
         "The intake queue prevents the register from becoming a dumping ground. Not every concern deserves formal lifecycle tracking. The triage step ensures that what enters the register has been validated and assigned."),
        ("Audit trail as default",
         "Every status change, every update, every promotion from intake is timestamped and attributed. This is not an optional feature \u2014 it is the foundation. When a regulator asks \"what did you do about X, and when?\", the answer is already recorded."),
        ("Zero infrastructure burden",
         "The solution runs on Azure Static Web Apps (free tier), stores data in SharePoint Lists (already provisioned), and authenticates through Entra ID (already in place). There are no servers to maintain, no databases to back up, and no new vendor relationships."),
    ]

    for title, desc in principles:
        h2(doc, title)
        body(doc, desc)

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    #  TECHNICAL FOUNDATION (brief)
    # ═══════════════════════════════════════════════
    h1(doc, "Technical Foundation")
    lead(doc, "A brief overview of how the solution fits within the existing enterprise environment.")

    tech = doc.add_table(rows=7, cols=2)
    tech.alignment = WD_TABLE_ALIGNMENT.LEFT
    style_table_borders(tech, "D4E2EE")

    tech_items = [
        ("Platform", "Azure Static Web Apps (free tier) \u2014 serverless, globally distributed"),
        ("Authentication", "Microsoft Entra ID (Azure AD) via MSAL \u2014 corporate SSO with MFA"),
        ("Data Storage", "SharePoint Online Lists within the PwC M365 tenant \u2014 no external databases"),
        ("Frontend", "React 19 single-page application with responsive design"),
        ("Security", "HTTPS-only, session-scoped tokens, PKCE auth flow, Content Security Policy headers"),
        ("Languages", "English and French with instant toggle"),
        ("Offline Capability", "CSV export for offline analysis; demo mode for disconnected environments"),
    ]
    for ri, (label, value) in enumerate(tech_items):
        cell_l = tech.rows[ri].cells[0]
        cell_v = tech.rows[ri].cells[1]
        set_cell_shading(cell_l, "F8FAFC")
        set_cell_margins(cell_l, 70, 70, 120, 120)
        set_cell_margins(cell_v, 70, 70, 120, 120)
        pl = cell_l.paragraphs[0]
        run(pl, label, bold=True, size=10, color=NAVY_800)
        pv = cell_v.paragraphs[0]
        run(pv, value, size=10, color=GRAY_600)

    spacer = doc.add_paragraph(); sp(spacer, 0, 16)

    body(doc, "The solution was designed to operate entirely within the existing Microsoft ecosystem. There are no external APIs, no third-party data processors, and no components outside PwC's security perimeter. Data never leaves the M365 tenant.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    #  WHAT COMES NEXT
    # ═══════════════════════════════════════════════
    h1(doc, "What Comes Next")
    lead(doc, "The current version delivers the core lifecycle management workflow. Future enhancements are shaped by how the team uses the tool in practice.")

    roadmap = [
        ("Automated notifications",
         "Teams or email alerts when items cross staleness thresholds, giving owners a gentle nudge before the meeting."),
        ("Committee-level analytics",
         "Trend data showing how the team's issue portfolio is evolving over time \u2014 average time to resolution, staleness trends, priority distribution shifts."),
        ("Role-based views",
         "Custom default views for different personas: a committee chair sees the full dashboard; an issue owner sees their portfolio filtered automatically."),
        ("Integration with ServiceNow",
         "For issues that originate from formal incident management, a direct feed from ServiceNow into the intake queue eliminates manual re-entry."),
    ]

    for title, desc in roadmap:
        h2(doc, title)
        body(doc, desc)

    body(doc, "Each of these enhancements builds on the same foundation: a structured lifecycle, visible accountability, and a low-friction workflow that makes the right thing the easy thing.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════
    #  CLOSING
    # ═══════════════════════════════════════════════
    for _ in range(4):
        s = doc.add_paragraph(); sp(s, 0, 0)

    close_rule = doc.add_paragraph()
    close_rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(close_rule, "\u2500" * 30, size=11, color=PWC_ORANGE)
    sp(close_rule, 0, 24)

    close_t = doc.add_paragraph()
    close_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(close_t, 0, 16)
    run(close_t, "Issues do not resolve themselves.\nBut they move faster when everyone can see them.", size=14, color=NAVY_700, font="Calibri")

    close_rule2 = doc.add_paragraph()
    close_rule2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run(close_rule2, "\u2500" * 30, size=11, color=PWC_ORANGE)
    sp(close_rule2, 0, 40)

    f1 = doc.add_paragraph()
    f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(f1, 0, 6)
    run(f1, "One Firm Risk Tracker  \u2014  Solution Overview v1.0  \u2014  February 2026", size=10, color=GRAY_400)

    f2 = doc.add_paragraph()
    f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sp(f2, 0, 0)
    run(f2, "PwC Canada  \u2014  OFR Risk & Quality Team", bold=True, size=10, color=NAVY_700)

    # ── Save ──
    doc.save(OUTPUT)
    print(f"\nSolution Overview DOCX created: {OUTPUT}")
    print(f"File size: {os.path.getsize(OUTPUT):,} bytes")


if __name__ == "__main__":
    main()
